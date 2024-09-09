import requests
from docx import Document
import os
import csv
from openai import OpenAI
import time

# Instantiate OpenAI client
client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))

# Initialize a new Word document
document = Document()

# Your Readwise API token from environment variables
api_token = os.environ.get("READWISE")

# Set up the headers with your Readwise API token
headers = {"Authorization": f"Token {api_token}"}

# CSV file to save the data
csv_file = "Highlights_with_Tags.csv"


def generate_tags_from_openai(highlight_text):
    """
    Use OpenAI to generate up to 3 tags for a given highlight using the chat completion model.
    """
    try:
        prompt = f'Generate up to 3 high-level relevant tags for the following highlight: "{highlight_text}". They should be separated by commas and not have hashes. Use British English. They should be restricted to: Economics, Technology, Startups, Science, Physics, Biology, Chemistry, Entrepreneurship, Liberalism, Philosophy, Environment, Religion, Politics, History, Psychology, Sociology, Statistics, United Kingdom, Quotes, Film, Music, Marketing, Politics, Personal Finance, Design, CBT, Lifetips, Europe, United States, Critical Thinking, IdPol, Health, Finance, Agriculture, Productivity, and Literature.'
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a helpful assistant."},
                {"role": "user", "content": prompt},
            ],
            max_tokens=50,
            temperature=0.5,
        )
        # Extract the tags from the response content
        tags_text = response.choices[0].message.content.strip()
        tags = [tag.strip() for tag in tags_text.split(",") if tag.strip()]
        return tags
    except Exception as e:
        print(f"Error generating tags: {e}")
        return []



def fetch_highlights():
    url = "https://readwise.io/api/v2/export/"
    next_page_cursor = None
    all_highlights = []

    while True:
        # Add pagination parameter if there is a next page
        params = {"pageCursor": next_page_cursor} if next_page_cursor else {}
        response = requests.get(url, headers=headers, params=params)

        if response.status_code == 200:
            data_export = response.json()

            # Add results from this page to the highlights list
            all_highlights.extend(data_export.get("results", []))

            # Check if there's another page of data
            next_page_cursor = data_export.get("nextPageCursor")
            if not next_page_cursor:
                break  # No more pages, exit loop
        else:
            print(f"Failed to fetch export data. Status code: {response.status_code}")
            print(response.text)
            break

    return all_highlights

try:
    # Fetch all highlights across all pages
    all_books = fetch_highlights()

    # Open a CSV file to write data
    with open(csv_file, mode="w", newline="", encoding="utf-8") as file:
        writer = csv.writer(file)
        writer.writerow(["Book Title", "Highlight", "Tags"])  # CSV header

        # Iterate over each book
        for book in all_books:
            title = book.get("title", "Untitled")  # Use 'Untitled' as fallback if title is missing
            highlights = book.get("highlights", [])

            # Add the book title to the document
            para = document.add_paragraph()
            para.add_run(title).bold = True

            # Add each highlight to the document
            for highlight in highlights:
                highlight_text = highlight.get("text", "").strip()
                if highlight_text:
                    document.add_paragraph(highlight_text)

                    # Get the existing tags from the highlight
                    existing_tags = [tag["name"] for tag in highlight.get("tags", [])]

                    # If there are fewer than 3 tags, use OpenAI to generate additional tags
                    if len(existing_tags) < 3:
                        additional_tags = generate_tags_from_openai(highlight_text)
                        # Add only enough tags to make the total equal to 3
                        existing_tags.extend(additional_tags[: (3 - len(existing_tags))])

                    # Add the tags to the document
                    if existing_tags:
                        tags_text = ", ".join(existing_tags)
                        document.add_paragraph(f"Tags: {tags_text}").italic = True

                    # Write to CSV and print the entry in the terminal
                    writer.writerow([title, highlight_text, ", ".join(existing_tags)])
                    print(f"CSV entry logged: Title: {title}, Highlight: {highlight_text}, Tags: {', '.join(existing_tags)}")

    # Save the document
    document.save("Highlights_with_Tags.docx")
    print("Highlights have been successfully saved to 'Highlights_with_Tags.docx' and CSV file.")

except requests.exceptions.RequestException as e:
    # Handle any network or request errors
    print(f"Request error: {e}")

except KeyError as e:
    # Handle any KeyError that may arise if a key is missing
    print(f"Key error: {e}")

except Exception as e:
    # General error handler to catch any other exceptions
    print(f"An error occurred: {e}")

# CSV file with the highlights and tags
csv_file = "Highlights_with_Tags.csv"

# Log file to track successfully updated highlight IDs
success_log_file = "updated_highlights_log.txt"

# Maximum number of retries
MAX_RETRIES = 5

def load_updated_highlight_ids():
    """
    Load the IDs of highlights that have already been successfully updated from the log file.
    """
    if os.path.exists(success_log_file):
        with open(success_log_file, "r") as f:
            return {line.strip() for line in f.readlines()}
    return set()

def log_successful_update(highlight_id):
    """
    Log a successfully updated highlight ID to the file.
    """
    with open(success_log_file, "a") as f:
        f.write(f"{highlight_id}\n")

def update_highlight_tags(highlight_id, tags):
    """
    Updates the tags for a specific highlight using Readwise API with exponential backoff.
    """
    url = f"https://readwise.io/api/v2/highlights/{highlight_id}/"
    payload = {"tags": [{"name": tag} for tag in tags]}
    retries = 0
    backoff = 2  # Start with 2 seconds

    while retries < MAX_RETRIES:
        try:
            # Send a PATCH request to update the highlight
            response = requests.patch(url, headers=headers, json=payload)

            if response.status_code == 200:
                print(f"Successfully updated highlight {highlight_id} with tags: {tags}")
                log_successful_update(highlight_id)  # Log the successful update
                return True
            elif response.status_code == 429:
                # Handle rate limiting
                print(f"Rate limited. Waiting for {backoff} seconds before retrying.")
                time.sleep(backoff)
                backoff *= 2  # Double the backoff time
                retries += 1
            else:
                # Handle other errors
                print(f"Failed to update highlight {highlight_id}. Status code: {response.status_code}")
                print(response.text)
                return False
        except requests.exceptions.RequestException as e:
            print(f"Request error: {e}")
            return False

    print(f"Max retries exceeded for highlight {highlight_id}.")
    return False


def fetch_highlights():
    """
    Fetch highlights from Readwise and return them.
    """
    try:
        response = requests.get("https://readwise.io/api/v2/export/", headers=headers)
        if response.status_code == 200:
            return response.json()["results"]
        else:
            print(f"Failed to fetch highlights. Status code: {response.status_code}")
            print(response.text)
            return []
    except requests.exceptions.RequestException as e:
        print(f"Request error: {e}")
        return []


def update_tags_from_csv():
    """
    Reads the CSV file and updates the corresponding highlights with new tags.
    """
    highlights_data = fetch_highlights()

    # Load the IDs of already updated highlights from the log file
    updated_highlight_ids = load_updated_highlight_ids()

    with open(csv_file, mode="r", encoding="utf-8") as file:
        reader = csv.DictReader(file)
        for row in reader:
            title = row["Book Title"]
            highlight_text = row["Highlight"]
            tags = row["Tags"].split(", ")

            matching_highlight = None
            for book in highlights_data:
                if book["title"] == title:
                    for highlight in book["highlights"]:
                        if highlight["text"] == highlight_text:
                            matching_highlight = highlight
                            break

            if matching_highlight:
                highlight_id = matching_highlight["id"]

                # Skip the highlight if it has already been updated
                if str(highlight_id) in updated_highlight_ids:
                    print(f"Skipping already updated highlight {highlight_id}.")
                    continue

                # Update the highlight's tags
                update_highlight_tags(highlight_id, tags)
            else:
                print(f"No matching highlight found for: {highlight_text}")


# Start the process
update_tags_from_csv()